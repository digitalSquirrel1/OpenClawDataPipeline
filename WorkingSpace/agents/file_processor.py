# -*- coding: utf-8 -*-
"""
Step 3 — File Processor  (parallel edition)
创建 spec 中所有文件，使用 ThreadPoolExecutor 并行处理。

多媒体文件策略（music, images, video）：
  sub_type=mp3/wav/flac/aac  → 强制下载，失败则跳过（不降级）
  sub_type=jpg/png/gif/svg    → 强制下载，失败则跳过（不降级）
  sub_type=mp4/mkv/avi/mov   → 强制下载，失败则跳过（不降级）
  特点：只使用下载策略，不降级到LLM生成或Jina抓取

下载策略（downloadable）：
  sub_type=pdf   → search_for_pdf() → download_binary() → 真实 .pdf
  sub_type=html  → search() → download_html() → 真实 .html
  sub_type=excel → search_for_filetype() → download_binary() → 真实 .xlsx/.xls
  sub_type=csv   → search_for_filetype() → download_binary() → .csv
  任何 sub_type → 若以上失败 → Jina 文本抓取 → .md 兜底

生成策略（generated）：
  csv/xlsx → LLM 生成 CSV 数据（openpyxl 转 xlsx）
  docx     → python-docx 生成（或纯文本兜底）
  md/txt   → LLM 生成 Markdown/文本
  pptx     → LLM 生成结构化大纲文本（保存为 .txt）
  py/json  → LLM 生成代码/JSON

并发控制：
  max_workers  — 并行文件数（默认 8，I/O密集型用线程池）
  _llm_sem     — LLM 并发上限（默认 4），避免 API 限速
  _search_sem  — Serper 搜索并发上限（默认 5）

多媒体文件特点：
  • 强制使用下载策略

  • 下载失败直接跳过，不降级处理

  • 增加下载尝试次数（6次）

  • 支持较大超时时间（视频120秒，音乐60秒）

  • 文件大小验证（至少1KB）

"""
import csv, io, json, os, time, re
import threading
import concurrent.futures
from pathlib import Path
from utils.llm_client import LLMClient
from utils.web_tools   import WebTools

GEN_SYSTEM = "你是一名专业的内容生成助手，能够生成高度仿真的职场文件内容，语言专业、数据真实。"

# 多媒体文件类型（强制下载，不降级到生成策略）
MEDIA_EXTENSIONS = {
    # 音乐
    ".mp3", ".wav", ".flac", ".aac", ".ogg", ".wma", ".m4a",
    # 图片
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".svg", ".tiff", ".tif",
    # 视频
    ".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm", ".m4v", ".3gp",
    "audio", "image", "video"  # 通用媒体类型标识
}

MEDIA_TYPES = {
    # 音乐
    "mp3", "wav", "flac", "aac", "ogg", "wma", "m4a", "audio",
    # 图片
    "jpg", "jpeg", "png", "gif", "bmp", "webp", "svg", "tiff", "tif", "image",
    # 视频
    "mp4", "avi", "mkv", "mov", "wmv", "flv", "webm", "m4v", "3gp", "video"
}

# 默认并发参数（可通过环境变量覆盖）
DEFAULT_WORKERS    = int(os.getenv("MAX_WORKERS",    "8"))
DEFAULT_LLM_SEM    = int(os.getenv("MAX_LLM_CALLS",  "4"))
DEFAULT_SEARCH_SEM = int(os.getenv("MAX_SEARCHES",   "5"))


class FileProcessor:
    def __init__(self, llm: LLMClient, web: WebTools, output_dir: str,
                 max_workers: int = DEFAULT_WORKERS):
        self.llm        = llm
        self.web        = web
        self.output_dir = Path(output_dir)
        self.max_workers = max_workers

        # 并发控制
        self._print_lock  = threading.Lock()
        self._llm_sem     = threading.Semaphore(DEFAULT_LLM_SEM)
        self._search_sem  = threading.Semaphore(DEFAULT_SEARCH_SEM)

    # ── thread-safe logging ────────────────────────────────────────────────
    def _log(self, msg: str) -> None:
        with self._print_lock:
            print(msg)

    # ── rate-limited wrappers ──────────────────────────────────────────────
    def _llm_generate(self, prompt, system=GEN_SYSTEM,
                      temperature=0.7, max_tokens=3000) -> str:
        with self._llm_sem:
            return self.llm.generate(prompt, system=system,
                                     temperature=temperature,
                                     max_tokens=max_tokens)

    def _llm_generate_json(self, prompt, max_tokens=3000) -> dict:
        with self._llm_sem:
            return self.llm.generate_json(prompt)

    def _search(self, query, num=5) -> list:
        with self._search_sem:
            return self.web.search(query, num=num)

    def _search_for_pdf(self, query, num=8) -> list:
        with self._search_sem:
            return self.web.search_for_pdf(query, num=num)

    def _search_for_filetype(self, query, exts, num=6) -> list:
        with self._search_sem:
            return self.web.search_for_filetype(query, exts=exts, num=num)

    def _search_and_fetch(self, query, num=4):
        with self._search_sem:
            return self.web.search_and_fetch(query, num=num)

    def _is_media_file(self, sub_type: str, file_path: Path) -> bool:
        """判断是否为多媒体文件（音乐、图片、视频）"""
        sub_lower = sub_type.lower()
        ext_lower = file_path.suffix.lower()

        # 检查sub_type
        if sub_lower in MEDIA_TYPES:
            return True

        # 检查文件扩展名
        if ext_lower in MEDIA_EXTENSIONS:
            return True

        return False

    # ──────────────────────────────────────────────────────────────────────────
    # Media files handler (music, images, videos) - download only, no fallback
    # ──────────────────────────────────────────────────────────────────────────
    def _handle_media_download(self, fspec: dict, abs_path: Path) -> bool:
        """处理多媒体文件下载（音乐、图片、视频）- 强制下载，失败则跳过"""
        sub_type = fspec.get("sub_type", "").lower()
        file_ext = abs_path.suffix.lower()

        # 确定文件类型标识
        media_type = sub_type if sub_type in MEDIA_TYPES else file_ext.lower().lstrip(".")

        self._log(f"    [MEDIA] 检测到{media_type}多媒体文件，尝试下载...")

        query = fspec.get("search_query", fspec.get("description", ""))

        # 针对不同类型使用不同的搜索策略
        if media_type in ("mp3", "wav", "flac", "aac", "ogg", "wma", "m4a", "audio"):
            # 音乐文件搜索
            search_query = f"{query} filetype:{media_type}" if media_type != "audio" else f"{query} music audio"
        elif media_type in ("jpg", "jpeg", "png", "gif", "bmp", "webp", "svg", "tiff", "tif", "image"):
            # 图片文件搜索
            search_query = f"{query} filetype:{media_type}" if media_type != "image" else f"{query} image photo"
        else:
            # 视频文件搜索
            search_query = f"{query} filetype:{media_type}" if media_type != "video" else f"{query} video"

        # 执行搜索和下载尝试
        download_attempts = 6  # 多媒体文件多尝试几次
        for attempt in range(download_attempts):
            try:
                results = self._search(search_query, num=4)
                if not results:
                    continue

                for r in results:
                    url = r.get("link", "")
                    if not url:
                        continue

                    self._log(f"    [MEDIA-{attempt+1}] 尝试下载: {url[:75]}...")

                    # 增加超时时间，多媒体文件通常较大
                    timeout = 120 if media_type in ("video", "mp4", "mkv", "avi") else 60
                    ok = self.web.download_binary(url, str(abs_path), timeout=timeout)

                    if ok and abs_path.exists():
                        file_size = abs_path.stat().st_size

                        # 检查文件大小是否合理
                        min_size = 1000  # 至少1KB
                        if file_size < min_size:
                            self._log(f"    [SKIP] 文件过小 ({file_size} bytes)，可能下载失败")
                            continue

                        size_str = self._format_file_size(file_size)
                        self._log(f"    [OK] {media_type.upper()} {size_str} -> {abs_path.name}")
                        return True

                    time.sleep(0.3)  # 下载失败后稍作等待

            except Exception as e:
                self._log(f"    [Error] 下载尝试失败: {e}")
                continue

        # 所有下载尝试都失败
        self._log(f"    [FAIL] {media_type.upper()} 文件下载失败，已跳过（不降级到生成策略）")
        return False

    def _format_file_size(self, size_bytes: int) -> str:
        """格式化文件大小显示"""
        if size_bytes >= 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.2f} MB"
        elif size_bytes >= 1024:
            return f"{size_bytes / 1024:.2f} KB"
        else:
            return f"{size_bytes} bytes"

    # ──────────────────────────────────────────────────────────────────────────
    # Public entry
    # ──────────────────────────────────────────────────────────────────────────
    def process(self, spec: dict, profile: dict) -> list[str]:
        base  = self.output_dir / "computer_profile"
        files = spec.get("files", [])
        total = len(files)

        # 创建所有目录（单线程，避免竞争）
        for d in spec.get("directories", []):
            (base / d).mkdir(parents=True, exist_ok=True)

        print(f"[Step 3] 并行处理 {total} 个文件（max_workers={self.max_workers}）...")

        done_count = [0]
        count_lock = threading.Lock()
        created: list[str] = []
        created_lock = threading.Lock()

        def process_one(idx_fspec):
            i, fspec = idx_fspec
            path     = fspec.get("path", "")
            ftype    = fspec.get("type", "generated")
            sub_type = fspec.get("sub_type", "")
            tag      = f"{ftype[:3].upper()}/{sub_type}" if sub_type else ftype[:3].upper()
            abs_path = base / path
            abs_path.parent.mkdir(parents=True, exist_ok=True)

            self._log(f"  [{i:02d}/{total}] [{tag}] {path}")

            # 检查是否为多媒体文件（音乐、图片、视频）
            is_media = self._is_media_file(sub_type, abs_path)

            try:
                if is_media:
                    # 多媒体文件：强制使用下载策略，失败则跳过
                    ok = self._handle_media_download(fspec, abs_path)
                elif ftype == "downloadable":
                    ok = self._handle_downloadable(fspec, abs_path, profile)
                else:
                    ok = self._handle_generated(fspec, abs_path, profile)
            except Exception as exc:
                self._log(f"    [Error] {exc}")
                ok = False

            with count_lock:
                done_count[0] += 1
                n = done_count[0]

            if ok:
                result = str(abs_path)
                if not abs_path.exists():
                    for ext in (".md", ".csv", ".txt"):
                        candidate = abs_path.with_suffix(ext)
                        if candidate.exists():
                            result = str(candidate)
                            break
                with created_lock:
                    created.append(result)
                self._log(f"    -> [OK] ({n}/{total} 完成)")
            else:
                self._log(f"    -> [FAIL] ({n}/{total} 完成)")

        with concurrent.futures.ThreadPoolExecutor(
                max_workers=self.max_workers,
                thread_name_prefix="fp") as executor:
            list(executor.map(process_one, enumerate(files, 1)))

        print(f"  -> 成功创建 {len(created)}/{total} 个文件")
        return created

    # ──────────────────────────────────────────────────────────────────────────
    # Downloadable dispatcher
    # ──────────────────────────────────────────────────────────────────────────
    def _handle_downloadable(self, fspec: dict, abs_path: Path,
                             profile: dict) -> bool:
        sub = fspec.get("sub_type", "").lower()
        ext = abs_path.suffix.lower()
        if sub == "pdf" or ext == ".pdf":
            return self._dl_pdf(fspec, abs_path, profile)
        elif sub == "html" or ext in (".html", ".htm", ".mhtml"):
            return self._dl_html(fspec, abs_path, profile)
        elif sub in ("excel", "xlsx", "xls") or ext in (".xlsx", ".xls"):
            return self._dl_excel(fspec, abs_path, profile)
        elif sub == "csv" or ext == ".csv":
            return self._dl_csv_web(fspec, abs_path, profile)
        else:
            return self._dl_generic(fspec, abs_path, profile)

    # ── PDF ───────────────────────────────────────────────────────────────────
    def _dl_pdf(self, fspec, abs_path, profile):
        query    = fspec.get("search_query", fspec.get("description", ""))
        pdf_urls = self._search_for_pdf(query, num=8)
        for url in pdf_urls:
            self._log(f"    [PDF] {url[:75]}")
            ok = self.web.download_binary(url, str(abs_path), timeout=60)
            if ok:
                kb = abs_path.stat().st_size // 1024
                self._log(f"    [OK] PDF {kb} KB -> {abs_path.name}")
                return True
            time.sleep(0.2)
        self._log(f"    [!] PDF 直链失败，降级 Jina...")
        return self._jina_fallback(fspec, abs_path, profile)

    # ── HTML page ─────────────────────────────────────────────────────────────
    def _dl_html(self, fspec, abs_path, profile):
        query   = fspec.get("search_query", fspec.get("description", ""))
        results = self._search(query, num=5)
        for r in results:
            url = r.get("link", "")
            if not url:
                continue
            self._log(f"    [HTML] {url[:75]}")
            html = self.web.download_html(url)
            if html and len(html) > 500:
                html_path = abs_path.with_suffix(".html")
                html_path.write_bytes(html)
                kb = html_path.stat().st_size // 1024
                self._log(f"    [OK] HTML {kb} KB -> {html_path.name}")
                return True
            time.sleep(0.2)
        return self._jina_fallback(fspec, abs_path.with_suffix(".html"), profile)

    # ── Excel ─────────────────────────────────────────────────────────────────
    def _dl_excel(self, fspec, abs_path, profile):
        query   = fspec.get("search_query", fspec.get("description", ""))
        xl_urls = self._search_for_filetype(query, exts=["xlsx", "xls"], num=6)
        for url in xl_urls:
            self._log(f"    [XLS] {url[:75]}")
            ok = self.web.download_binary(url, str(abs_path), timeout=45)
            if ok:
                kb = abs_path.stat().st_size // 1024
                self._log(f"    [OK] Excel {kb} KB -> {abs_path.name}")
                return True
            time.sleep(0.2)
        self._log(f"    [!] Excel 下载失败，LLM 生成...")
        prompt = fspec.get("content_prompt", fspec.get("description", ""))
        return self._gen_xlsx(prompt, abs_path, profile)

    # ── CSV from web ──────────────────────────────────────────────────────────
    def _dl_csv_web(self, fspec, abs_path, profile):
        query    = fspec.get("search_query", fspec.get("description", ""))
        csv_urls = self._search_for_filetype(query, exts=["csv"], num=4)
        for url in csv_urls:
            self._log(f"    [CSV] {url[:75]}")
            ok = self.web.download_binary(url, str(abs_path), timeout=30)
            if ok:
                kb = abs_path.stat().st_size // 1024
                self._log(f"    [OK] CSV {kb} KB -> {abs_path.name}")
                return True
        self._log(f"    [!] CSV 下载失败，LLM 生成...")
        prompt = fspec.get("content_prompt", fspec.get("description", ""))
        return self._gen_csv(prompt, abs_path, profile)

    # ── Generic downloadable ──────────────────────────────────────────────────
    def _dl_generic(self, fspec, abs_path, profile):
        query   = fspec.get("search_query", fspec.get("description", ""))
        results = self._search(query, num=4)
        for r in results:
            url = r.get("link", "")
            if not url:
                continue
            ok = self.web.download_binary(url, str(abs_path), timeout=30)
            if ok:
                return True
        return self._jina_fallback(fspec, abs_path.with_suffix(".md"), profile)

    # ── Jina text fallback ────────────────────────────────────────────────────
    def _jina_fallback(self, fspec, path, profile):
        query = fspec.get("search_query", fspec.get("description", ""))
        src_url, content = self._search_and_fetch(query, num=4)
        if content:
            desc = fspec.get("description", "")
            ext  = path.suffix.lower()
            if ext == ".html":
                html_content = (
                    f"<!DOCTYPE html><html><head>"
                    f"<meta charset='utf-8'><title>{desc}</title></head><body>"
                    f"<h1>{desc}</h1>"
                    f"<p><em>来源: <a href='{src_url}'>{src_url}</a></em></p>"
                    f"<pre>{content}</pre></body></html>"
                )
                path.write_text(html_content, encoding="utf-8")
            else:
                path = path.with_suffix(".md")
                header = f"# {desc}\n\n> 来源: {src_url}\n\n"
                path.write_text(header + content, encoding="utf-8")
            kb = path.stat().st_size // 1024
            self._log(f"    [OK] Jina {kb} KB -> {path.name}")
            return True
        self._log(f"    [!] Jina 失败，LLM 兜底...")
        return self._llm_fallback(fspec, path.with_suffix(".md"), profile)

    # ──────────────────────────────────────────────────────────────────────────
    # Generated files dispatcher
    # ──────────────────────────────────────────────────────────────────────────
    def _handle_generated(self, fspec, abs_path, profile):
        fmt    = fspec.get("format", _infer_fmt(abs_path))
        prompt = fspec.get("content_prompt", fspec.get("description", ""))
        dispatch = {
            "csv":    self._gen_csv,
            "xlsx":   self._gen_xlsx,
            "docx":   self._gen_docx,
            "pptx":   self._gen_pptx,
            "py":     self._gen_code,
            "python": self._gen_code,
            "json":   self._gen_json_file,
        }
        fn = dispatch.get(fmt, self._gen_text)
        return fn(prompt, abs_path, profile)

    # ── Generators ────────────────────────────────────────────────────────────
    def _gen_text(self, prompt, path, profile):
        ctx     = f"你的角色：{profile.get('role')}，单位：{profile.get('company')}\n\n"
        content = self._llm_generate(
            ctx + prompt + "\n\n请直接输出文件正文，不要附加任何说明。",
            temperature=0.8, max_tokens=3000
        )
        path.write_text(content, encoding="utf-8")
        self._log(f"    [OK] 文本 {len(content)} chars -> {path.name}")
        return True

    def _gen_csv(self, prompt, path, profile):
        ctx = f"用户角色：{profile.get('role')}，单位：{profile.get('company')}\n\n"
        raw = self._llm_generate(
            ctx + prompt + "\n\n只输出CSV内容（第一行列名），不要使用markdown代码块，不要附加任何说明。",
            temperature=0.3, max_tokens=3000
        )
        raw = _strip_fences(raw)
        path.write_text(raw, encoding="utf-8-sig")
        self._log(f"    [OK] CSV {raw.count(chr(10))} 行 -> {path.name}")
        return True

    def _gen_xlsx(self, prompt, path, profile):
        csv_path = path.with_suffix(".csv")
        self._gen_csv(prompt, csv_path, profile)
        try:
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            with open(csv_path, encoding="utf-8-sig") as f:
                for row in csv.reader(f):
                    ws.append(row)
            for col in ws.columns:
                max_len = max((len(str(c.value or "")) for c in col), default=8)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 40)
            wb.save(str(path))
            csv_path.unlink(missing_ok=True)
            kb = path.stat().st_size // 1024
            self._log(f"    [OK] XLSX {kb} KB -> {path.name}")
        except ImportError:
            csv_path.rename(path.with_suffix(".csv"))
        return True

    def _gen_docx(self, prompt, path, profile):
        ctx     = f"用户角色：{profile.get('role')}，单位：{profile.get('company')}\n\n"
        content = self._llm_generate(
            ctx + prompt + "\n\n请以Markdown格式输出文档正文（标题用#，段落分明），不要附加任何说明。",
            temperature=0.75, max_tokens=3000
        )
        try:
            from docx import Document
            doc = Document()
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
            doc.save(str(path))
            kb = path.stat().st_size // 1024
            self._log(f"    [OK] DOCX {kb} KB -> {path.name}")
        except ImportError:
            md_path = path.with_suffix(".md")
            md_path.write_text(content, encoding="utf-8")
            self._log(f"    [OK] DOCX->MD {len(content)} chars -> {md_path.name}")
        return True

    def _gen_pptx(self, prompt, path, profile):
        ctx     = f"用户角色：{profile.get('role')}，单位：{profile.get('company')}\n\n"
        content = self._llm_generate(
            ctx + prompt +
            "\n\n请生成完整的PPT演讲稿大纲（逐页列出页码、标题、要点、备注），专业风格，不少于15页。",
            temperature=0.75, max_tokens=3000
        )
        txt_path = path.with_suffix(".txt")
        txt_path.write_text(content, encoding="utf-8")
        self._log(f"    [OK] PPT大纲 {len(content)} chars -> {txt_path.name}")
        return True

    def _gen_code(self, prompt, path, profile):
        code = self._llm_generate(
            prompt + "\n\n只输出Python代码，不要有任何额外说明。",
            temperature=0.4, max_tokens=2000
        )
        path.write_text(_strip_fences(code), encoding="utf-8")
        self._log(f"    [OK] Python {len(code)} chars -> {path.name}")
        return True

    def _gen_json_file(self, prompt, path, profile):
        data = self._llm_generate_json(prompt)
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self._log(f"    [OK] JSON -> {path.name}")
        return True

    def _llm_fallback(self, fspec, path, profile):
        desc   = fspec.get("description", "")
        prompt = (
            f"请仿真生成以下文档内容（Markdown格式，不少于800字，内容专业真实）：\n{desc}\n\n"
            f"用户背景：{profile.get('role')}，{profile.get('company')}"
        )
        content = self._llm_generate(prompt, temperature=0.7, max_tokens=2000)
        header  = f"# {desc}\n\n> 注: 由AI仿真生成（原始文件未能在线获取）\n\n"
        path.write_text(header + content, encoding="utf-8")
        self._log(f"    [OK] LLM兜底 {len(content)} chars -> {path.name}")
        return True


# ── helpers ───────────────────────────────────────────────────────────────────
def _infer_fmt(path: Path) -> str:
    return path.suffix.lower().lstrip(".") or "txt"


def _strip_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        end   = len(lines) - 1 if lines[-1].strip() == "```" else len(lines)
        text  = "\n".join(lines[1:end])
    return text